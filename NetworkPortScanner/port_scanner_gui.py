import socket
import threading
import tkinter as tk
from tkinter import messagebox, filedialog
import logging
import sys
from fpdf import FPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont

# Global variable to control scan stopping
stop_scan = False

# Set up logging
logging.basicConfig(filename="port_scanner.log", level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Function to generate a report in PDF format
def generate_pdf_report(target, start_port, end_port, scan_results):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Port Scan Report", ln=True, align='C')
    pdf.cell(200, 10, txt=f"Target: {target}", ln=True, align='L')
    pdf.cell(200, 10, txt=f"Ports scanned: {start_port} to {end_port}", ln=True, align='L')
    pdf.multi_cell(200, 10, txt="Scan Results:\n\n" + scan_results)
    file_name = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
    if file_name:
        pdf.output(file_name)
        logging.info(f"PDF report generated: {file_name}")
        messagebox.showinfo("Success", "PDF Report Generated Successfully")

# Function to generate a report in Word format
def generate_word_report(target, start_port, end_port, scan_results):
    doc = Document()
    doc.add_heading('Port Scan Report', 0)
    doc.add_paragraph(f"Target: {target}")
    doc.add_paragraph(f"Ports scanned: {start_port} to {end_port}")
    doc.add_paragraph("Scan Results:")
    doc.add_paragraph(scan_results)
    file_name = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if file_name:
        doc.save(file_name)
        logging.info(f"Word report generated: {file_name}")
        messagebox.showinfo("Success", "Word Report Generated Successfully")

# Function to generate a report in Image format
def generate_image_report(target, start_port, end_port, scan_results):
    img = Image.new('RGB', (600, 400), color='white')
    d = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    d.text((10, 10), f"Port Scan Report\nTarget: {target}\nPorts scanned: {start_port} to {end_port}\n\nScan Results:\n\n{scan_results}", fill='black', font=font)
    file_name = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png")])
    if file_name:
        img.save(file_name)
        logging.info(f"Image report generated: {file_name}")
        messagebox.showinfo("Success", "Image Report Generated Successfully")

# Function to scan ports
def scan_ports(target, start_port, end_port, output_file, scan_results=None):
    global stop_scan  # Use the global stop flag
    if scan_results is not None:
        scan_results.delete(1.0, tk.END)  # Clear previous results in GUI
    logging.info(f"Scanning ports from {start_port} to {end_port} on {target}")
    
    for port in range(start_port, end_port + 1):
        if stop_scan:
            if scan_results is not None:
                scan_results.insert(tk.END, "Scan Stopped\n")
            logging.info("Scan stopped by user")
            return
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(1)
            result = sock.connect_ex((target, port))
            if result == 0:
                if scan_results is not None:
                    scan_results.insert(tk.END, f"Port {port}: OPEN\n")
                logging.info(f"Port {port} is OPEN")
            else:
                if scan_results is not None:
                    scan_results.insert(tk.END, f"Port {port}: CLOSED\n")
                logging.info(f"Port {port} is CLOSED")
            sock.close()
        except socket.error as e:
            if scan_results is not None:
                scan_results.insert(tk.END, f"Port {port}: FILTERED\n")
            logging.error(f"Error scanning port {port}: {e}")
    
    if scan_results is not None:
        scan_results.insert(tk.END, "Scan Complete\n")
    logging.info("Scan complete")
    
    if output_file:
        generate_report(target, start_port, end_port, scan_results.get(1.0, tk.END) if scan_results is not None else "No Results")

# Function to generate the report (calls different report generators based on choice)
def generate_report(target, start_port, end_port, scan_results):
    result = messagebox.askquestion("Report Type", "Which format do you want to export the report to?", icon='question')
    if result == 'yes':
        generate_pdf_report(target, start_port, end_port, scan_results)
    elif result == 'no':
        result_word = messagebox.askquestion("Word Report", "Would you like to generate a Word report?", icon='question')
        if result_word == 'yes':
            generate_word_report(target, start_port, end_port, scan_results)
    else:
        generate_image_report(target, start_port, end_port, scan_results)

# Function to start scanning from the GUI
def start_scan_gui():
    global stop_scan
    target = entry_target.get()
    start_port = int(entry_start_port.get())
    end_port = int(entry_end_port.get())
    output_file = var_export_report.get()

    if not target or start_port < 1 or end_port > 65535:
        messagebox.showerror("Input Error", "Please provide valid input")
        return

    stop_scan = False
    thread = threading.Thread(target=scan_ports, args=(target, start_port, end_port, output_file, scan_results))
    thread.start()

# Function to stop scanning
def stop_scan_func():
    global stop_scan
    stop_scan = True
    logging.info("Scan manually stopped by user")

# Function to create the GUI
def create_gui():
    global scan_results, entry_target, entry_start_port, entry_end_port, var_export_report
    root = tk.Tk()
    root.title("Network Port Scanner")

    # Target IP or domain
    tk.Label(root, text="Target (IP or Domain)").pack()
    entry_target = tk.Entry(root)
    entry_target.pack()

    # Start Port
    tk.Label(root, text="Start Port").pack()
    entry_start_port = tk.Entry(root)
    entry_start_port.pack()

    # End Port
    tk.Label(root, text="End Port").pack()
    entry_end_port = tk.Entry(root)
    entry_end_port.pack()

    # Export Report
    tk.Label(root, text="Do you want to export report?").pack()
    var_export_report = tk.StringVar(value="no")
    tk.Radiobutton(root, text="Yes", variable=var_export_report, value="yes").pack()
    tk.Radiobutton(root, text="No", variable=var_export_report, value="no").pack()

    # Start Scan Button
    tk.Button(root, text="Start Scan", command=start_scan_gui).pack()

    # Stop Scan Button
    tk.Button(root, text="Stop Scan", command=stop_scan_func).pack()

    # Results display area
    scan_results = tk.Text(root, height=10, width=50)
    scan_results.pack()

    root.mainloop()

if __name__ == "__main__":
    # Check if arguments are passed for command-line execution
    if len(sys.argv) > 1:
        target = sys.argv[1]
        start_port = int(sys.argv[2])
        end_port = int(sys.argv[3])
        output_file = sys.argv[4] if len(sys.argv) > 4 else None
        scan_ports(target, start_port, end_port, output_file)
    else:
        create_gui()
